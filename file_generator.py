
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color
from reportlab.lib.units import inch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, RGBColor

class FileGenerator:
    def createPDF(id, forename, surname, category):
        #Creating file
        canvas = Canvas(f"reconocimiento-{id}.pdf", pagesize=letter)
        width, height = letter

        #Background image and rectangle for information
        bg_im = ImageReader("assets/bg_image.jpg")
        logo = ImageReader("assets/logo.png")
        canvas.drawImage(image=bg_im, x=0, y=0, width=width, height=height)
        canvas.drawImage(image=logo, x=inch/2, y=height-1.5*inch, width=inch, height=inch)
        gray50transparent = Color(160, 160, 160, alpha=0.75)
        canvas.setFillColor(gray50transparent)
        canvas.rect(0.5*inch, 2*inch, width-inch, height-4*inch, fill=True, stroke=False)

        #Text
        canvas.setFont('Helvetica', 20)
        canvas.setFillColor(Color(0,0,0, alpha=1)) #Black

        canvas.drawCentredString(width/2, 8*inch, "Torneo de Programación Competitiva")
        canvas.drawCentredString(width/2, 7.5*inch, "Copa Guadalajara 2021")
        canvas.drawCentredString(width/2, 6*inch, "Se otorga el reconocimiento a:")
        canvas.setFont('Helvetica-Bold', 20)
        canvas.drawCentredString(width/2, 5.5*inch, f'{forename} {surname}')
        canvas.setFont('Helvetica', 20)
        canvas.drawCentredString(width/2, 4*inch, f"Por su participación en la categoría {category}")
        canvas.drawCentredString(width/2, 3.5*inch, 'en la copa de programación competitiva')
        canvas.drawCentredString(width/2, 3*inch, 'Guadalajara 2021.')
        
        #Closing file
        canvas.save()

    def createTXT(id, forename, surname, category):
        f = open(f"reconocimiento-{id}.txt", 'w')
        f.write(f"Torneo de Programación Competitiva\nCopa Guadalajara 2021\n\nSe otorga el reconocimiento a:\n{forename} {surname}")
        f.write(f'\n\nPor su participación en la categoría {category}\nen la copa de programación competitiva\nGuadalajara 2021.')
        f.close()

    def createDOCX(id, forename, surname, category):
        f = Document()

        normal = f.styles['Normal']
        font = normal.font
        font.name = 'Helvetica'
        font.size = Pt(14)
        
        styles = f.styles
        my_style1 = styles.add_style('MyStyle1', WD_STYLE_TYPE.CHARACTER)
        font = my_style1.font
        font.name = 'Helvetica'
        font.size = Pt(18)

        styles = f.styles
        my_style2 = styles.add_style('MyStyle2', WD_STYLE_TYPE.CHARACTER)
        font = my_style2.font
        font.name = 'Helvetica'
        font.size = Pt(24)
        font.color.rgb = RGBColor(255, 153, 51)

        heading = f.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = heading.add_run("Torneo de Programación Competitiva\nCopa Guadalajara 2021", style='MyStyle2').bold = True

        paragraph1 = f.add_paragraph(f"\n\n\nSe otorga el reconocimiento a:\n")
        paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph2 = f.add_paragraph()
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = paragraph2.add_run(f'{forename} {surname}\n', style='MyStyle1').bold = True

        paragraph3 = f.add_paragraph()
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = paragraph3.add_run()
        r3.add_text(f'\n\nPor su participación en la categoría {category}\nen la copa de programación competitiva\nGuadalajara 2021.')

        f.add_paragraph()
        f.add_paragraph()
        f.add_paragraph()
        paragraph4 = f.add_paragraph()
        paragraph4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = paragraph4.add_run()
        r4.add_picture("assets/logo.png", width=Inches(3), height=Inches(3))


        f.save(f"reconocimiento-{id}.docx")

    def createTicket(id, data):
        f = open(f"Ticket-{id}.txt", 'w')
        f.write("COMPROBANTE DE PAGO DEL PARTICIPANTE\n\n")

        fields = ["Nombre" , "Apellido", "CURP", "Sexo", "Estado", "Ciudad", "Colonia", "Calle", "Número", "C.P.", "Estudiante", "Escuela" ,"Categoría"]
        for i, field in enumerate(fields):
            f.write(f"{field}: {data[i]}\n")
        
        if data[-1] == 'Infantil':
            cost = 29.95
        elif data[-1] == 'Aficionados':
            cost = 58.99
        elif data[-1] == 'Avanzado':
            cost = 89.85
        else:
            cost = 49.95

        f.write(f"\nTotal: ${cost}")
        f.close()